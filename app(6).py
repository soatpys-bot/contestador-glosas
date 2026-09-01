import io
import json
import os
import re
import subprocess
import tempfile
from datetime import date
from pathlib import Path

import streamlit as st
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image
from google import genai
from google.genai import types

APP_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = APP_DIR / "PLANTILLA OFICIAL GLOSAS.docx"
NORM_DIR = APP_DIR / "normativa"

CODE_CONTEXT = {
    "RE9601": "Opción disponible para revisión humana.",
    "RE9602": "Opción predeterminada para revisión humana.",
    "RE9701": "Opción disponible para revisión humana.",
    "RE9801": "Opción disponible para revisión humana.",
    "RE9901": "Opción disponible para revisión humana.",
}

st.set_page_config(page_title="Contestador Inteligente de Glosas", page_icon="📄", layout="wide")
st.title("📄 Contestador Inteligente de Glosas")
st.caption("Análisis dentro de la aplicación con Gemini · revisión humana · respuesta institucional · PDF consolidado")

def clean(v):
    return "" if v is None else str(v).strip()

def pdf_text(data):
    reader=PdfReader(io.BytesIO(data))
    return "\n".join((p.extract_text() or "") for p in reader.pages).strip()

def docx_text(data):
    d=Document(io.BytesIO(data))
    out=[p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out).strip()

def read_any(name,data):
    ext=Path(name).suffix.lower()
    if ext==".pdf": return pdf_text(data)
    if ext==".docx": return docx_text(data)
    if ext==".txt": return data.decode("utf-8",errors="ignore")
    return ""

def replace_everywhere(doc, mapping):
    # Paragraphs
    for p in doc.paragraphs:
        for old,new in mapping.items():
            if old in p.text:
                # Preserve paragraph style while replacing runs.
                full=p.text
                for run in p.runs:
                    run.text=""
                p.runs[0].text=full.replace(old,new) if p.runs else p.add_run(full.replace(old,new))
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old,new in mapping.items():
                        if old in p.text:
                            full=p.text
                            for run in p.runs: run.text=""
                            p.runs[0].text=full.replace(old,new) if p.runs else p.add_run(full.replace(old,new))
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            for p in cell2.paragraphs:
                                for old,new in mapping.items():
                                    if old in p.text:
                                        full=p.text
                                        for run in p.runs: run.text=""
                                        p.runs[0].text=full.replace(old,new) if p.runs else p.add_run(full.replace(old,new))

def set_cell_text(cell, value):
    """Replace a table cell's visible text without leaving template sample data."""
    value = clean(value)
    # Keep the cell's paragraph/style; remove all old runs so sample values cannot survive.
    p = cell.paragraphs[0]
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = value
    else:
        p.add_run(value)
    # Remove any extra paragraphs from the cell, which could otherwise retain example text.
    for extra in list(cell.paragraphs[1:]):
        parent = extra._element.getparent()
        if parent is not None:
            parent.remove(extra._element)


def replace_paragraph_exact(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = clean(new)
            else:
                p.add_run(clean(new))
            return True
    return False


def fill_template(meta, justification, out_docx):
    """Fill the supplied institutional template using its known table structure.

    The values visible in the original template are examples only. Every variable
    field is overwritten before the document is saved; generation is blocked when
    required real data is missing so an example cannot accidentally be published.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError("No se encontró PLANTILLA OFICIAL GLOSAS.docx en el proyecto.")

    required = {
        "entidad": meta.get("entidad"),
        "factura": meta.get("factura"),
        "valor_glosado": meta.get("valor_glosado"),
        "codigo_respuesta": meta.get("codigo_respuesta"),
        "valor_aceptado": meta.get("valor_aceptado"),
        "valor_objetado": meta.get("valor_objetado"),
        "fecha": meta.get("fecha"),
        "justificacion": justification,
    }
    missing = [k for k, v in required.items() if not clean(v)]
    if missing:
        raise ValueError("Faltan campos obligatorios antes de generar el PDF: " + ", ".join(missing))

    doc = Document(str(TEMPLATE_PATH))

    # Exact structure of the institutional template supplied by the user:
    # table 0: title/date; table 1: invoice/value; table 2: response detail.
    if len(doc.tables) < 3:
        raise ValueError("La plantilla institucional no tiene la estructura esperada de 3 tablas.")

    t0 = doc.tables[0]
    if len(t0.rows) < 1 or len(t0.columns) < 3:
        raise ValueError("La tabla de encabezado de la plantilla no tiene la estructura esperada.")
    set_cell_text(t0.cell(0, 2), f"FECHA\n{meta['fecha']}")

    t1 = doc.tables[1]
    if len(t1.rows) < 2 or len(t1.columns) < 2:
        raise ValueError("La tabla de factura/valor de la plantilla no tiene la estructura esperada.")
    set_cell_text(t1.cell(0, 1), meta["factura"])
    set_cell_text(t1.cell(1, 1), meta["valor_glosado"])

    t2 = doc.tables[2]
    if len(t2.rows) < 2 or len(t2.columns) < 4:
        raise ValueError("La tabla de respuesta de la plantilla no tiene la estructura esperada.")
    set_cell_text(t2.cell(1, 0), meta["codigo_respuesta"])
    set_cell_text(t2.cell(1, 1), meta["valor_aceptado"])
    set_cell_text(t2.cell(1, 2), meta["valor_objetado"])
    set_cell_text(t2.cell(1, 3), justification)

    # Entity is a paragraph in the supplied template.
    if not replace_paragraph_exact(doc, "PREVISORA SOAT", meta["entidad"]):
        raise ValueError("No se encontró el campo de entidad 'PREVISORA SOAT' en la plantilla suministrada.")

    doc.save(str(out_docx))

    # Safety check: none of the known sample values may remain anywhere in the
    # generated document's visible text.
    check = Document(str(out_docx))
    visible = "\n".join(p.text for p in check.paragraphs)
    for t in check.tables:
        for row in t.rows:
            visible += "\n" + " | ".join(c.text for c in row.cells)
    sample_values = ["PREVISORA SOAT", "FEDV348712", "398.000", "RE998", "398000"]
    leaked = [x for x in sample_values if x in visible]
    if leaked:
        raise ValueError("La plantilla conserva datos de ejemplo que no fueron reemplazados: " + ", ".join(leaked))
def libreoffice_convert(input_path, output_dir):
    commands=[
        ["libreoffice","--headless","--convert-to","pdf","--outdir",str(output_dir),str(input_path)],
        ["soffice","--headless","--convert-to","pdf","--outdir",str(output_dir),str(input_path)],
    ]
    last=None
    for cmd in commands:
        try:
            r=subprocess.run(cmd,stdout=subprocess.PIPE,stderr=subprocess.PIPE,text=True,timeout=120)
            if r.returncode==0:
                pdf=Path(output_dir)/(Path(input_path).stem+".pdf")
                if pdf.exists(): return pdf
            last=r.stderr or r.stdout
        except Exception as e: last=str(e)
    raise RuntimeError("No se pudo convertir el documento institucional a PDF. Detalle: "+clean(last))

def file_to_pdf(name,data,tmp):
    ext=Path(name).suffix.lower()
    if ext==".pdf":
        p=Path(tmp)/(Path(name).stem+".pdf"); p.write_bytes(data); return p
    if ext==".docx":
        src=Path(tmp)/(Path(name).stem+".docx"); src.write_bytes(data)
        return libreoffice_convert(src,tmp)
    if ext in {".png",".jpg",".jpeg"}:
        im=Image.open(io.BytesIO(data)).convert("RGB")
        p=Path(tmp)/(Path(name).stem+".pdf"); im.save(p,"PDF",resolution=150); return p
    raise ValueError("Formato no soportado: "+name)

def merge_pdfs(paths):
    """Une varios PDF usando la API compatible con pypdf 5.x+."""
    writer = PdfWriter()
    for p in paths:
        writer.append(str(p))
    out = io.BytesIO()
    writer.write(out)
    writer.close()
    return out.getvalue()

def get_key():
    try:
        k=clean(st.secrets.get("GEMINI_API_KEY",""))
        if k: return k
    except Exception: pass
    return clean(os.getenv("GEMINI_API_KEY",""))

def gemini_analyze(api_key, model, glosa_bytes, glosa_text, supports_text, norm_text, user_entity):
    client=genai.Client(api_key=api_key)
    schema={
      "type":"object",
      "properties":{
        "entidad":{"type":"string"},
        "factura":{"type":"string"},
        "valor_glosado":{"type":"string"},
        "codigo_respuesta_sugerido":{"type":"string"},
        "argumentacion":{"type":"string"},
        "soportes_a_revisar":{"type":"array","items":{"type":"string"}},
        "observaciones":{"type":"string"}
      },
      "required":["entidad","factura","valor_glosado","codigo_respuesta_sugerido","argumentacion","soportes_a_revisar","observaciones"]
    }
    prompt=f"""
Analiza esta glosa para preparar una respuesta institucional de una IPS en Colombia.

REGLAS:
- Lee el PDF completo y, si existe, usa también su membrete/logo para identificar la entidad.
- Si la entidad no puede determinarse con seguridad, usa el dato proporcionado por el usuario.
- No inventes factura, valores, hechos, procedimientos, diagnósticos, fechas, soportes ni autorizaciones.
- No busques ni inventes códigos de glosa. Solo identifica la situación descrita.
- El código de respuesta por defecto de la aplicación es RE9602; el usuario podrá modificarlo después.
- Los datos visibles en la plantilla institucional son ejemplos y NUNCA deben copiarse como datos de la glosa actual.
- Debe existir UNA SOLA argumentación general para la factura, aunque tenga varias glosas.
- Si no hay soporte adicional, igual redacta una respuesta argumentativa basada únicamente en la información disponible.
- Si hay soportes, intégralos al razonamiento sin inventar su contenido.
- Usa la normativa suministrada cuando sea pertinente. No cites una norma que no esté disponible.
- La respuesta debe ser formal, técnica, clara y firme, pero no afirmar hechos que no estén soportados.
- No incluyas el significado del código dentro de la argumentación.
- El resultado será revisado por una persona antes de radicarse.

DATO DE ENTIDAD PROPORCIONADO POR USUARIO (puede estar vacío): {user_entity}

TEXTO EXTRAÍDO DEL PDF:
{glosa_text[:70000]}

TEXTO DE SOPORTES:
{supports_text[:70000]}

NORMATIVA:
{norm_text[:90000]}

Devuelve JSON.
"""
    parts=[
        types.Part.from_bytes(data=glosa_bytes,mime_type="application/pdf"),
        prompt
    ]
    resp=client.models.generate_content(
        model=model,
        contents=parts,
        config=types.GenerateContentConfig(
            temperature=0.1,
            response_mime_type="application/json",
            response_schema=schema
        )
    )
    return json.loads(clean(resp.text))

def normative_text():
    chunks=[]
    for p in NORM_DIR.rglob("*") if NORM_DIR.exists() else []:
        if p.is_file() and p.suffix.lower() in {".pdf",".docx",".txt"}:
            try:
                chunks.append(f"===== {p.name} =====\n{read_any(p.name,p.read_bytes())[:120000]}")
            except Exception as e:
                chunks.append(f"===== {p.name} =====\n[No se pudo leer: {e}]")
    return "\n".join(chunks)

# Sidebar
with st.sidebar:
    st.header("Configuración")
    key_manual=st.text_input("Clave Gemini (no la compartas)",type="password")
    model=st.selectbox("Modelo Gemini",["gemini-2.5-flash-lite","gemini-2.5-flash"],index=0)
    st.divider()
    st.markdown("### Códigos de respuesta")
    for c,desc in CODE_CONTEXT.items():
        st.write(f"**{c}** — {desc}")
    st.caption("El contexto se muestra aquí para ayudarte a elegir; NO se imprime en el PDF.")
    if key_manual:
        st.session_state["gemini_key"]=key_manual

api_key=clean(st.session_state.get("gemini_key","")) or get_key()
if api_key: st.success("Gemini configurado.")
else: st.warning("Configura GEMINI_API_KEY en Streamlit Secrets o pega la clave temporalmente en la barra lateral.")

st.subheader("1. Cargar glosa")
glosa=st.file_uploader("PDF de la glosa",type=["pdf"],accept_multiple_files=False)

st.subheader("2. Cargar soportes adicionales (opcional y múltiples)")
supports=st.file_uploader("Soportes PDF, DOCX, PNG o JPG",type=["pdf","docx","png","jpg","jpeg"],accept_multiple_files=True)

if glosa:
    gb=glosa.getvalue()
    try: gt=pdf_text(gb)
    except Exception as e:
        st.error(f"No se pudo leer el PDF: {e}"); st.stop()

    st.subheader("3. Datos que puedes revisar antes del análisis")
    entity_hint=st.text_input("Entidad a la que se responderá (si no se identifica claramente, escríbela aquí)")
    with st.expander("Texto extraído de la glosa",expanded=False):
        st.text_area("Texto",gt[:30000],height=260,disabled=True)
    if supports:
        with st.expander(f"Soportes cargados: {len(supports)}",expanded=False):
            for f in supports: st.write("•",f.name)

    if st.button("🔎 Analizar glosa con Gemini",type="primary",use_container_width=True):
        if not api_key:
            st.error("Falta la clave Gemini.")
            st.stop()
        support_chunks=[]
        for f in supports or []:
            try:
                support_chunks.append(f"===== {f.name} =====\n{read_any(f.name,f.getvalue())[:60000]}")
            except Exception as e:
                support_chunks.append(f"===== {f.name} =====\n[No se pudo leer: {e}]")
        with st.spinner("Analizando PDF, soportes y normativa..."):
            try:
                res=gemini_analyze(api_key,model,gb,gt,"\n".join(support_chunks),normative_text(),entity_hint)
                st.session_state["analysis"]=res
                st.session_state["support_files"]=[(f.name,f.getvalue(),f.type) for f in supports or []]
                st.success("Análisis terminado. Revisa todo antes de generar el PDF.")
            except Exception as e:
                st.error(f"Error al consultar Gemini: {e}")

if "analysis" in st.session_state:
    a=st.session_state["analysis"]
    st.subheader("4. Revisión completa antes de generar el PDF")

    c1,c2=st.columns(2)
    entidad=c1.text_input("Entidad",a.get("entidad",""),key="final_entidad")
    factura=c2.text_input("N. de factura",a.get("factura",""),key="final_factura")
    valor_glosado=c1.text_input("Valor glosado",a.get("valor_glosado",""),key="final_valor_glosado")
    codigo=c2.selectbox("Código de respuesta",list(CODE_CONTEXT.keys()),index=list(CODE_CONTEXT.keys()).index("RE9602"),key="final_codigo")
    st.info(f"Contexto del código seleccionado: {CODE_CONTEXT[codigo]}  ·  Este contexto NO aparecerá en el PDF.")

    aceptado=st.text_input("Valor aceptado","0",key="final_aceptado")
    objetado=st.text_input("Valor objetado",value=a.get("valor_glosado",""),key="final_objetado")

    just=st.text_area("JUSTIFICACIÓN — una sola respuesta general para la factura",a.get("argumentacion",""),height=380,key="final_just")
    obs=st.text_area("Observaciones internas (NO se imprime en el PDF)",a.get("observaciones",""),height=120,key="final_obs")

    st.markdown("**Soportes que Gemini recomienda revisar/anexar**")
    rec="\n".join(a.get("soportes_a_revisar",[]))
    st.text_area("Lista de soportes",rec,height=120,key="final_support_list")

    meta={
        "fecha":date.today().strftime("%d/%m/%Y"),
        "entidad":entidad,
        "factura":factura,
        "valor_glosado":valor_glosado,
        "codigo_respuesta":codigo,
        "valor_aceptado":aceptado,
        "valor_objetado":objetado,
    }

    st.warning("La plantilla contiene datos de ejemplo. Antes de generar, verifica entidad, factura, valor, código, valores aceptado/objetado y la justificación. La aplicación bloquea la generación si falta un campo obligatorio o si detecta datos de ejemplo en el documento final.")

    st.subheader("5. Generar PDF institucional + soportes")
    st.caption("La fecha se genera con el día en que se produce la respuesta.")

    if st.button("📄 Generar PDF final",type="primary",use_container_width=True):
        with st.spinner("Generando la plantilla institucional y anexando soportes..."):
            try:
                with tempfile.TemporaryDirectory() as td:
                    td=Path(td)
                    out_docx=td/"Respuesta_institucional.docx"
                    fill_template(meta,just,out_docx)
                    institutional_pdf=libreoffice_convert(out_docx,td)

                    pdf_paths=[institutional_pdf]
                    for name,data,mime in st.session_state.get("support_files",[]):
                        pdf_paths.append(file_to_pdf(name,data,td))

                    final_pdf=merge_pdfs(pdf_paths)
                    st.session_state["final_pdf"]=final_pdf
                    st.session_state["final_filename"]=f"Respuesta_Glosas_{factura or 'sin_factura'}.pdf"
                st.success("PDF final generado correctamente.")
            except Exception as e:
                st.error(f"No se pudo generar el PDF final: {e}")

    if st.session_state.get("final_pdf"):
        st.download_button("⬇️ Descargar PDF final",data=st.session_state["final_pdf"],
                           file_name=st.session_state["final_filename"],mime="application/pdf",
                           use_container_width=True)
