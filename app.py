import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date
from pathlib import Path

import streamlit as st
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image
from google import genai
from google.genai import types

APP_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = APP_DIR / "PLANTILLA OFICIAL GLOSAS.docx"
NORM_DIR = APP_DIR / "normativa"

DEFAULT_CODE = "RE9602"
CODE_CONTEXT = {
    "RE9601": "El prestador aporta evidencia que demuestra que la devolución es injustificada al 100%.",
    "RE9602": "El prestador aporta evidencia que demuestra que la glosa es injustificada al 100%.",
    "RE9701": "El prestador informa que la devolución ha sido aceptada al 100%.",
    "RE9702": "El prestador informa que la glosa ha sido aceptada al 100%.",
    "RE9801": "El prestador informa que la glosa ha sido aceptada y subsanada parcialmente.",
    "RE9901": "El prestador informa que la glosa, siendo justificada, ha podido ser subsanada totalmente.",
}

st.set_page_config(page_title="Contestador Inteligente de Glosas", page_icon="📄", layout="wide")

st.markdown("""
<style>
    .main-title {
        font-size: 2.25rem;
        font-weight: 800;
        margin-bottom: 0.15rem;
    }
    .main-subtitle {
        color: #667085;
        font-size: 1rem;
        margin-bottom: 1rem;
    }
    .step-card {
        border: 1px solid #E4E7EC;
        border-radius: 14px;
        padding: 0.9rem 1rem;
        background: #FFFFFF;
        margin: 0.35rem 0 0.8rem 0;
    }
    .step-title {
        font-size: 1.05rem;
        font-weight: 750;
        margin-bottom: 0.2rem;
    }
    .step-help {
        color: #667085;
        font-size: 0.9rem;
    }
    .status-box {
        border: 1px solid #D0D5DD;
        border-radius: 12px;
        padding: 0.7rem 0.9rem;
        background: #F8FAFC;
        margin: 0.5rem 0 1rem 0;
    }
    .defense-box {
        border-left: 5px solid #344054;
        border-radius: 10px;
        padding: 0.85rem 1rem;
        background: #F9FAFB;
        margin: 0.7rem 0;
    }
    .final-box {
        border: 2px solid #D0D5DD;
        border-radius: 14px;
        padding: 1rem;
        background: #FCFCFD;
        margin-top: 1rem;
    }
    div.stButton > button {
        border-radius: 10px;
        font-weight: 700;
        min-height: 2.7rem;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">📄 Contestador Inteligente de Glosas</div>', unsafe_allow_html=True)
st.markdown('<div class="main-subtitle">Análisis de glosa · defensa de la IPS · revisión humana · previsualización · PDF institucional + anexos</div>', unsafe_allow_html=True)

# Identificador del caso actual. Cambiarlo también reinicia los widgets de carga.
if "case_id" not in st.session_state:
    st.session_state["case_id"] = 0

# Reinicio seguro del caso actual. No modifica archivos del proyecto ni la configuración de Gemini.
if st.button("🆕 NUEVA GLOSA", use_container_width=True):
    keys_to_clear = [
        "analysis", "support_files", "preview_pdf", "final_pdf", "final_filename",
        "final_entidad", "final_factura", "final_valor_glosado", "final_codigo",
        "final_eval", "final_suggestion", "reviewer_decision", "final_aceptado",
        "final_objetado", "final_just", "final_obs", "final_support_list",
        "glosa", "supports", "entity_hint",
    ]
    for key in keys_to_clear:
        st.session_state.pop(key, None)
    st.session_state["case_id"] += 1
    st.rerun()

st.markdown("""
<div class="status-box">
<b>Flujo de trabajo:</b>
① Ingresar glosa &nbsp;→&nbsp; ② Soportes &nbsp;→&nbsp; ③ Análisis &nbsp;→&nbsp;
④ Revisión humana &nbsp;→&nbsp; ⑤ Previsualización &nbsp;→&nbsp; ⑥ Aprobar y generar PDF
</div>
""", unsafe_allow_html=True)



def clean(value):
    return "" if value is None else str(value).strip()


def pdf_text(data):
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def docx_text(data):
    doc = Document(io.BytesIO(data))
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            out.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(out).strip()


def read_any(name, data):
    ext = Path(name).suffix.lower()
    if ext == ".pdf":
        return pdf_text(data)
    if ext == ".docx":
        return docx_text(data)
    if ext == ".txt":
        return data.decode("utf-8", errors="ignore")
    return ""


@st.cache_data(show_spinner=False)
def load_normativa():
    chunks = []
    if not NORM_DIR.exists():
        return ""
    for path in sorted(NORM_DIR.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in {".pdf", ".docx", ".txt"}:
            continue
        try:
            chunks.append(f"===== {path.name} =====\n{read_any(path.name, path.read_bytes())[:140000]}")
        except Exception as exc:
            chunks.append(f"===== {path.name} =====\n[No se pudo leer: {exc}]")
    return "\n\n".join(chunks)


def get_key():
    try:
        key = clean(st.secrets.get("GEMINI_API_KEY", ""))
        if key:
            return key
    except Exception:
        pass
    return clean(os.getenv("GEMINI_API_KEY", ""))


def clear_cell(cell, value):
    """Replace all visible cell text, retaining the template cell formatting."""
    value = clean(value)
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = value
    else:
        p.add_run(value)
    for extra in list(cell.paragraphs[1:]):
        parent = extra._element.getparent()
        if parent is not None:
            parent.remove(extra._element)


def replace_paragraph_exact(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = clean(new)
            else:
                p.add_run(clean(new))
            return True
    return False


def visible_doc_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def fill_template(meta, justification, out_docx):
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError("No se encontró PLANTILLA OFICIAL GLOSAS.docx en el proyecto.")

    required = ["entidad", "factura", "valor_glosado", "codigo_respuesta", "valor_aceptado", "valor_objetado", "fecha"]
    missing = [key for key in required if not clean(meta.get(key))]
    if not clean(justification):
        missing.append("justificacion")
    if missing:
        raise ValueError("Faltan campos obligatorios: " + ", ".join(missing))

    doc = Document(str(TEMPLATE_PATH))
    if len(doc.tables) != 3:
        raise ValueError("La plantilla oficial debe contener exactamente 3 tablas.")

    t0, t1, t2 = doc.tables
    clear_cell(t0.cell(0, 2), f"FECHA\n{meta['fecha']}")
    clear_cell(t1.cell(0, 1), meta["factura"])
    clear_cell(t1.cell(1, 1), meta["valor_glosado"])
    clear_cell(t2.cell(1, 0), meta["codigo_respuesta"])
    clear_cell(t2.cell(1, 1), meta["valor_aceptado"])
    clear_cell(t2.cell(1, 2), meta["valor_objetado"])
    clear_cell(t2.cell(1, 3), justification)

    # La versión actual de la plantilla limpia contiene "EPS o Aseguradora" como texto normal.
    # También aceptamos la versión anterior que tenía un marcador.
    if not replace_paragraph_exact(doc, "EPS o Aseguradora", meta["entidad"]):
        replace_paragraph_exact(doc, "{{eps o aseguradora}}", meta["entidad"])

    doc.save(str(out_docx))

    # Seguridad: ningún marcador ni dato de ejemplo puede pasar al documento generado.
    check = Document(str(out_docx))
    visible = visible_doc_text(check)
    forbidden = [
        "{{eps o aseguradora}}", "{{fecha}}", "{{factura}}", "{{valor_glosado}}",
        "{{codigo_respuesta}}", "{{valor_aceptado}}", "{{valor_objetado}}", "{{justificacion_ia}}",
        "PREVISORA SOAT", "FEDV348712", "398.000", "RE998", "398000",
    ]
    leaked = [item for item in forbidden if item.lower() in visible.lower()]
    if leaked:
        raise ValueError("Se detectaron marcadores/datos de ejemplo en el documento final: " + ", ".join(leaked))


def libreoffice_convert(input_path, output_dir):
    last_error = ""
    for binary in ("libreoffice", "soffice"):
        if not shutil.which(binary):
            continue
        try:
            result = subprocess.run(
                [binary, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(input_path)],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=180,
            )
            pdf_path = Path(output_dir) / (Path(input_path).stem + ".pdf")
            if result.returncode == 0 and pdf_path.exists():
                return pdf_path
            last_error = result.stderr or result.stdout or "Conversión sin archivo de salida."
        except Exception as exc:
            last_error = str(exc)
    raise RuntimeError("No se pudo convertir la plantilla a PDF. Verifica LibreOffice/packages.txt. " + clean(last_error))


def file_to_pdf(name, data, temp_dir):
    ext = Path(name).suffix.lower()
    target = Path(temp_dir) / (Path(name).stem + ".pdf")
    if ext == ".pdf":
        target.write_bytes(data)
        return target
    if ext == ".docx":
        source = Path(temp_dir) / (Path(name).stem + ".docx")
        source.write_bytes(data)
        return libreoffice_convert(source, temp_dir)
    if ext in {".png", ".jpg", ".jpeg"}:
        with Image.open(io.BytesIO(data)) as image:
            image.convert("RGB").save(target, "PDF", resolution=150.0)
        return target
    raise ValueError(f"Formato no soportado: {name}")


def merge_pdfs(paths):
    writer = PdfWriter()
    for path in paths:
        reader = PdfReader(str(path))
        for page in reader.pages:
            writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def render_pdf_preview(pdf_bytes):
    try:
        import fitz
    except Exception:
        return []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    try:
        for i in range(min(len(doc), 5)):
            pix = doc.load_page(i).get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
            images.append(pix.tobytes("png"))
    finally:
        doc.close()
    return images


def support_parts(supports):
    text_chunks = []
    media_parts = []
    for uploaded in supports or []:
        data = uploaded.getvalue()
        ext = Path(uploaded.name).suffix.lower()
        if ext in {".pdf", ".docx", ".txt"}:
            try:
                text_chunks.append(f"===== {uploaded.name} =====\n{read_any(uploaded.name, data)[:70000]}")
            except Exception as exc:
                text_chunks.append(f"===== {uploaded.name} =====\n[No se pudo leer: {exc}]")
        elif ext in {".png", ".jpg", ".jpeg"}:
            mime = "image/png" if ext == ".png" else "image/jpeg"
            media_parts.append(types.Part.from_bytes(data=data, mime_type=mime))
    return "\n\n".join(text_chunks), media_parts


def gemini_analyze(api_key, model, source_type, glosa_bytes, glosa_text, support_text, support_media, norm_text, user_entity):
    client = genai.Client(api_key=api_key)
    schema = {
        "type": "object",
        "properties": {
            "entidad": {"type": "string"},
            "factura": {"type": "string"},
            "valor_glosado": {"type": "string"},
            "evaluacion_glosa": {"type": "string"},
            "argumentacion": {"type": "string"},
            "sugerencia_revisor": {"type": "string"},
            "soportes_a_revisar": {"type": "array", "items": {"type": "string"}},
            "observaciones": {"type": "string"},
        },
        "required": ["entidad", "factura", "valor_glosado", "evaluacion_glosa", "argumentacion", "sugerencia_revisor", "soportes_a_revisar", "observaciones"],
    }

    source_instruction = "La glosa fue cargada como PDF." if source_type == "PDF" else "La glosa fue pegada/escrita directamente por el usuario en el cuadro de texto. No existe PDF de glosa; trabaja con ese texto como fuente principal."
    prompt = f"""
Eres el motor de análisis de un contestador institucional de glosas médicas en Colombia y trabajas para la IPS/prestador que debe responder una glosa formulada por una entidad responsable de pago.

OBJETIVO PRINCIPAL:
La finalidad de la respuesta es DEFENDER A LA IPS y CONTROVERTIR LA GLOSA cuando existan fundamentos documentales y normativos para hacerlo. No debes redactar como si representaras a la entidad pagadora ni limitarte a sostener la objeción.

REGLAS OBLIGATORIAS:
1. Analiza primero qué está objetando la entidad y cuáles son sus fundamentos.
2. Busca en la glosa, soportes y normativa elementos que permitan demostrar que la glosa NO procede o que procede solo parcialmente.
3. La argumentación debe ser firme, técnica y respetuosa, pero nunca inventar hechos.
4. Si falta un soporte que podría fortalecer la defensa, indícalo como soporte a revisar; no afirmes que existe.
5. Si, después de revisar la evidencia disponible, encuentras que un aspecto de la objeción parece procedente, NO decidas aceptarlo automáticamente. Expónlo únicamente como "sugerencia para el revisor" para que la persona decida si lo tiene en cuenta.
6. La decisión final siempre corresponde al revisor humano.
7. No inventes diagnósticos, procedimientos, fechas, valores, autorizaciones, soportes, contratos, tarifas, códigos de glosa, números de factura ni hechos clínicos.
8. No inventes códigos de respuesta. El código se selecciona manualmente en la aplicación; RE9602 aparece por defecto.
9. No copies ningún dato de ejemplo de la plantilla institucional.
10. Debe existir UNA SOLA argumentación general para la factura, aunque la glosa contenga varios motivos.
11. Usa únicamente la información realmente disponible en la glosa, los soportes y la normativa suministrada.
12. Si un dato esencial no aparece, déjalo vacío o señálalo en observaciones.
13. La descripción del código de respuesta NO debe incorporarse a la justificación.
14. La respuesta será revisada y editada por una persona antes de generar el PDF final.

{source_instruction}

ENTIDAD DADA POR EL USUARIO (puede estar vacía):
{user_entity}

TEXTO DE LA GLOSA:
{glosa_text[:100000]}

TEXTO DE SOPORTES:
{support_text[:100000]}

NORMATIVA DISPONIBLE:
{norm_text[:140000]}

En "evaluacion_glosa" explica brevemente si, con la evidencia disponible, la objeción parece injustificada, parcialmente sustentada o sustentada, y por qué.
En "argumentacion" redacta la defensa institucional de la IPS lista para ser revisada y editada. No redactes una aceptación salvo que el revisor decida posteriormente tenerla en cuenta.
En "sugerencia_revisor" indica, solo si aplica, algún aspecto que el revisor debería considerar antes de cerrar la respuesta. No reemplaces la decisión humana.
Devuelve exclusivamente JSON según el esquema solicitado.
"""

    contents = []
    if glosa_bytes:
        contents.append(types.Part.from_bytes(data=glosa_bytes, mime_type="application/pdf"))
    contents.append(prompt)
    contents.extend(support_media)

    response = client.models.generate_content(
        model=model,
        contents=contents,
        config=types.GenerateContentConfig(
            temperature=0.1,
            response_mime_type="application/json",
            response_schema=schema,
        ),
    )
    raw = clean(response.text)
    if not raw:
        raise RuntimeError("Gemini no devolvió una respuesta.")
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError("Gemini devolvió un formato que no pudo interpretarse como JSON.") from exc


with st.sidebar:
    st.header("⚙️ Configuración")
    key_manual = st.text_input("Clave Gemini (no la compartas)", type="password")
    model = st.selectbox("Modelo Gemini", ["gemini-3.5-flash-lite", "gemini-3.5-flash"], index=0)
    st.divider()
    st.markdown("### 📋 Código de respuesta")
    st.caption("RE9602 queda seleccionado por defecto. Puedes cambiarlo durante la revisión humana.")
    st.caption("El significado se muestra en pantalla para orientar la selección y nunca se imprime en el PDF.")

api_key = clean(key_manual) or get_key()
if api_key:
    st.success("Gemini configurado.")
else:
    st.warning("Escribe la clave Gemini en la barra lateral o configura GEMINI_API_KEY en Streamlit Secrets.")

norm_text = load_normativa()
if norm_text:
    st.sidebar.success("Normativa cargada dentro de la aplicación.")
else:
    st.sidebar.error("No se encontró la carpeta normativa/ con la Resolución 2284 de 2023.")

st.markdown('<div class="step-card"><div class="step-title">① Ingresar glosa</div><div class="step-help">Cargue el PDF o pegue/escriba directamente el texto de la glosa.</div></div>', unsafe_allow_html=True)
source_type = st.radio("¿Cómo tienes la glosa?", ["Cargar PDF", "Pegar o escribir texto"], horizontal=True, key=f"source_type_{st.session_state['case_id']}")

glosa = None
glosa_bytes = None
glosa_text = ""

if source_type == "Cargar PDF":
    glosa = st.file_uploader("Documento de glosa (PDF)", type=["pdf"], accept_multiple_files=False, key=f"glosa_{st.session_state['case_id']}")
    if glosa:
        glosa_bytes = glosa.getvalue()
        try:
            glosa_text = pdf_text(glosa_bytes)
        except Exception as exc:
            st.error(f"No se pudo leer el PDF de la glosa: {exc}")
            st.stop()
else:
    glosa_text = st.text_area(
        "Pega aquí el texto completo de la glosa",
        height=300,
        placeholder="Pega aquí la glosa, incluyendo entidad, factura, valores, motivos de objeción y cualquier detalle que haya comunicado la entidad responsable de pago.",
        help="No necesitas tener el PDF. Gemini analizará este texto como la fuente de la glosa.",
        key=f"glosa_text_{st.session_state['case_id']}",
    )

st.markdown('<div class="step-card"><div class="step-title">② Soportes adicionales</div><div class="step-help">Cargue los documentos que ayuden a sustentar la defensa de la IPS. Son opcionales y pueden ser varios.</div></div>', unsafe_allow_html=True)
supports = st.file_uploader(
    "PDF, DOCX o imágenes. Se incorporarán al PDF final como anexos.",
    type=["pdf", "docx", "png", "jpg", "jpeg"],
    accept_multiple_files=True,
    key=f"supports_{st.session_state['case_id']}",
)

has_glosa = bool(glosa_text.strip()) or bool(glosa_bytes)

if has_glosa:
    st.markdown('<div class="step-card"><div class="step-title">③ Análisis con Gemini</div><div class="step-help">Gemini propone una defensa con base en la glosa, los soportes y la normativa disponible.</div></div>', unsafe_allow_html=True)
    entity_hint = st.text_input("Entidad a la que se responderá (solo si no aparece claramente en la glosa)", key=f"entity_hint_{st.session_state['case_id']}")
    if glosa_text:
        with st.expander("Texto de la glosa", expanded=False):
            st.text_area("Contenido", glosa_text[:40000], height=260, disabled=True)
    if supports:
        with st.expander(f"Soportes cargados: {len(supports)}", expanded=False):
            for item in supports:
                st.write("•", item.name)

    if st.button("🔎 Analizar glosa con Gemini", type="primary", use_container_width=True):
        if not api_key:
            st.error("Falta la clave Gemini.")
            st.stop()
        with st.spinner("Gemini está analizando la glosa, buscando argumentos de defensa y revisando la normativa..."):
            try:
                support_text, media = support_parts(supports)
                result = gemini_analyze(api_key, model, source_type, glosa_bytes, glosa_text, support_text, media, norm_text, entity_hint)
                st.session_state["analysis"] = result
                st.session_state["support_files"] = [(item.name, item.getvalue(), item.type) for item in supports or []]
                st.session_state.pop("preview_pdf", None)
                st.session_state.pop("final_pdf", None)
                st.success("Análisis terminado. Revisa y corrige antes de producir el documento.")
            except Exception as exc:
                st.error(f"Error al consultar Gemini: {exc}")

if "analysis" in st.session_state:
    analysis = st.session_state["analysis"]
    st.markdown('<div class="step-card"><div class="step-title">④ Revisión humana obligatoria</div><div class="step-help">Gemini propone; usted revisa, corrige y decide qué respuesta se utilizará.</div></div>', unsafe_allow_html=True)
    st.markdown('<div class="defense-box"><b>🛡️ Regla principal:</b> Gemini propone argumentos para defender la posición de la IPS. <b>La decisión final siempre es humana.</b> La aplicación no radica nada automáticamente.</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    entidad = col1.text_input("Entidad", value=clean(analysis.get("entidad")), key="final_entidad")
    factura = col2.text_input("N. de factura", value=clean(analysis.get("factura")), key="final_factura")
    valor_glosado = col1.text_input("Valor glosado", value=clean(analysis.get("valor_glosado")), key="final_valor_glosado")
    codigo = col2.selectbox(
        "Código de respuesta",
        list(CODE_CONTEXT.keys()),
        index=list(CODE_CONTEXT).index(DEFAULT_CODE),
        key="final_codigo",
        format_func=lambda c: f"{c} — {CODE_CONTEXT[c]}",
    )
    st.info(f"**{codigo}**: {CODE_CONTEXT[codigo]}  ·  Este significado es informativo y NO se imprime en el PDF.")

    st.markdown("#### Evaluación de Gemini")
    st.text_area("¿Qué concluye Gemini sobre la procedencia de la glosa?", value=clean(analysis.get("evaluacion_glosa")), height=130, key="final_eval")
    st.text_area("Sugerencia de Gemini para el revisor (NO se imprime)", value=clean(analysis.get("sugerencia_revisor")), height=110, key="final_suggestion")
    decision = st.selectbox(
        "Decisión del revisor sobre la sugerencia de Gemini",
        [
            "Mantener la defensa de la IPS (no tomar la sugerencia)",
            "Tomar la sugerencia en cuenta parcialmente",
            "Tomar la sugerencia en cuenta",
        ],
        key="reviewer_decision",
    )

    aceptado = st.text_input("Valor aceptado", value="0", key="final_aceptado")
    objetado = st.text_input("Valor objetado", value=valor_glosado, key="final_objetado")
    justificacion = st.text_area(
        "JUSTIFICACIÓN — defensa de la IPS, una sola respuesta general para la factura",
        value=clean(analysis.get("argumentacion")),
        height=420,
        key="final_just",
        help="Edita aquí la respuesta definitiva. La orientación principal es defender la no aceptación de la glosa cuando la evidencia lo permita.",
    )
    st.text_area("Observaciones internas (NO se imprime en el PDF)", value=clean(analysis.get("observaciones")), height=120, key="final_obs")

    st.markdown("**Soportes que Gemini recomienda revisar/anexar**")
    st.text_area("Lista de soportes", value="\n".join(analysis.get("soportes_a_revisar", [])), height=120, key="final_support_list")

    st.markdown('<div class="step-card"><div class="step-title">⑤ Previsualización</div><div class="step-help">Revise el documento antes de aprobarlo. Esta copia todavía NO es el PDF definitivo.</div></div>', unsafe_allow_html=True)
    st.caption("Primero se genera una copia de revisión con la plantilla oficial. Nada se descarga como PDF final hasta que tú lo apruebes.")

    meta = {
        "fecha": date.today().strftime("%d/%m/%Y"),
        "entidad": entidad,
        "factura": factura,
        "valor_glosado": valor_glosado,
        "codigo_respuesta": codigo,
        "valor_aceptado": aceptado,
        "valor_objetado": objetado,
    }

    if st.button("👁️ PREVISUALIZAR PDF", type="primary", use_container_width=True):
        try:
            with tempfile.TemporaryDirectory() as temp:
                temp = Path(temp)
                out_docx = temp / "Respuesta_institucional.docx"
                fill_template(meta, justificacion, out_docx)
                preview_path = libreoffice_convert(out_docx, temp)
                st.session_state["preview_pdf"] = preview_path.read_bytes()
            st.session_state.pop("final_pdf", None)
            st.success("Previsualización generada. Revisa el documento antes de crear el PDF final.")
        except Exception as exc:
            st.error(f"No fue posible generar la previsualización: {exc}")

    if st.session_state.get("preview_pdf"):
        st.markdown("### 👁️ Revisa el documento antes de aprobarlo")
        preview_images = render_pdf_preview(st.session_state["preview_pdf"])
        if preview_images:
            for image in preview_images:
                st.image(image, use_container_width=True)
        else:
            st.info("La previsualización visual requiere PyMuPDF. Puedes continuar si el PDF de prueba se generó correctamente.")

        st.divider()
        st.markdown('<div class="final-box"><div class="step-title">⑥ Aprobar y generar PDF final</div><div class="step-help">Solo continúe después de revisar la previsualización y confirmar los datos.</div></div>', unsafe_allow_html=True)
        st.warning("Solo pulsa este botón después de comprobar que entidad, factura, valores, código y justificación están correctos. Los soportes cargados se anexarán después del documento institucional.")

        if st.button("✅ APROBAR Y GENERAR PDF FINAL", type="primary", use_container_width=True):
            try:
                with tempfile.TemporaryDirectory() as temp:
                    temp = Path(temp)
                    out_docx = temp / "Respuesta_institucional.docx"
                    fill_template(meta, justificacion, out_docx)
                    institutional_pdf = libreoffice_convert(out_docx, temp)
                    pdf_paths = [institutional_pdf]
                    for name, data, _mime in st.session_state.get("support_files", []):
                        pdf_paths.append(file_to_pdf(name, data, temp))
                    final_pdf = merge_pdfs(pdf_paths)
                st.session_state["final_pdf"] = final_pdf
                safe_invoice = re.sub(r"[^A-Za-z0-9_-]+", "_", factura or "sin_factura")
                st.session_state["final_filename"] = f"Respuesta_Glosas_{safe_invoice}.pdf"
                st.success("PDF final generado correctamente, con la respuesta institucional y los soportes anexos.")
            except Exception as exc:
                st.error(f"No se pudo generar el PDF final: {exc}")

    if st.session_state.get("final_pdf"):
        st.download_button(
            "⬇️ Descargar PDF final",
            data=st.session_state["final_pdf"],
            file_name=st.session_state["final_filename"],
            mime="application/pdf",
            use_container_width=True,
        )
else:
    st.markdown("""
### Flujo de trabajo
1. Elige **Cargar PDF** o **Pegar o escribir texto**.
2. Si no tienes PDF, pega la glosa completa en el cuadro.
3. Carga los soportes adicionales que tengas, si aplica.
4. Pulsa **Analizar glosa con Gemini**.
5. Gemini analizará la objeción con enfoque de **defensa de la IPS**.
6. Revisa entidad, factura, valor, código, valores aceptado/objetado y justificación.
7. Lee la sugerencia de Gemini; **tú decides** si la tienes en cuenta.
8. Pulsa **PREVISUALIZAR PDF**.
9. Revisa visualmente el documento.
10. Solo después pulsa **GENERAR PDF FINAL**.

**Importante:** la plantilla oficial es solo un formato. Sus datos de ejemplo nunca deben utilizarse como datos del caso y la aplicación bloquea la generación si detecta marcadores o ejemplos conocidos.
""")
